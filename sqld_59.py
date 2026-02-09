from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz, re, os, textwrap, shlex, subprocess
from collections import Counter

pdf_path="/mnt/data/59회기출변형.pdf"

# --- Extract stems/options ---
doc_pdf = fitz.open(pdf_path)
full_text = "\n".join([doc_pdf[i].get_text("text") for i in range(len(doc_pdf))])

pattern = re.compile(r'\n(\d{1,2})\.\s')
matches = list(pattern.finditer("\n"+full_text))
joined = "\n"+full_text
starts = [m.start() for m in matches]
qnums = [int(m.group(1)) for m in matches]

blocks = {}
for i, q in enumerate(qnums):
    s = starts[i]
    e = starts[i+1] if i+1 < len(starts) else len(joined)
    blocks[q] = joined[s:e].strip()

noise_exact = set([
    '출간 제안하기','도서 협찬 제안하기','자주 묻는 질문','도서','무료 원데이 특강','무료 저자 인강','시험 일정','CBT',
    '빠르게 따는 자격증','이벤트','아티클','교수회원 신청',
    '제공하는 콘텐츠 프로덕션 & 프로바이더 입니다. 골든래빗은 취미, 경제, 수험서, 만화, IT 등 다양한 분야에서 책을 제작하고 있습니다.',
    '골든','|','N'
])
def is_noise(s:str)->bool:
    if s in noise_exact:
        return True
    if s.startswith('남은 시간') or s.startswith('응답'):
        return True
    if '골든래빗' in s:
        return True
    return False

def parse_block(block:str):
    if "💡" in block:
        stem_part, opt_part = block.split("💡", 1)
        stem_part = stem_part.strip()
        opt_lines = [ln.strip() for ln in opt_part.strip().splitlines() if ln.strip()]
        opt_lines = [ln for ln in opt_lines if not is_noise(ln)]
    else:
        stem_part = block.strip()
        opt_lines = []
    m = re.match(r'^(\d{1,2})\.\s*(.*)$', stem_part, flags=re.S)
    stem = (m.group(2) if m else stem_part).strip()
    stem = re.sub(r'\s+', ' ', stem)
    return stem, opt_lines

qa = {}
for q in range(1, 51):
    stem, opts = parse_block(blocks[q])
    qa[q] = {"stem": stem, "opts": opts, "extras": []}

# --- Extras helpers ---
def add_table(q, title, headers, rows):
    qa[q]["extras"].append({"type":"table","title":title,"headers":headers,"rows":rows})

def add_code(q, title, code):
    qa[q]["extras"].append({"type":"code","title":title,"code":code})

def add_image(q, title, image_path, width_in=6.0):
    qa[q]["extras"].append({"type":"image","title":title,"path":image_path,"width_in":width_in})

def add_note(q, note):
    qa[q]["extras"].append({"type":"note","note":note})

# --- Add extras ---
add_table(7, "주문", ["주문ID","주문금액"], [["1","100"],["2","200"],["3","300"],["4","400"],["5","500"]])

add_table(9, "PRODUCT", ["ID","NAME","PRICE"], [["1","TV","1000"],["2","Laptop","2000"],["3","Tablet","NULL"],["4","Phone","1500"]])

erd_img="/mnt/data/crop_q10_erd.png"
if os.path.exists(erd_img):
    add_image(10, "ERD", erd_img, width_in=6.5)

add_note(11, "원본 PDF에서 11번 문항의 보기(가, 나) 문장이 포함된 구간이 누락되어, 아래처럼 확인 가능한 부분만 복원했습니다.")
add_code(11, "보기(부분)", "다. ROLE은 여러 사용자를 묶어서 관리하는 그룹 계정이다.\n라. ROLE은 다양한 권한들을 하나의 그룹으로 묶어 관리 편의성을 제공하는 객체이다.\n[가] (원본에서 확인 불가)\n[나] (원본에서 확인 불가)")

add_table(12, "매출", ["날짜","매출액"], [["2025-10-01","100"],["2025-10-02","200"],["2025-10-03","300"]])
add_table(12, "결과", ["날짜","매출액","PREV_SALES"], [["2025-10-03","300","200"],["2025-10-02","200","100"],["2025-10-01","100","NULL"]])
add_code(12, "SQL", "SELECT 날짜, 매출액,\n       ______(매출액) OVER (ORDER BY 날짜 DESC) AS PREV_SALES\nFROM 매출;")

add_note(14, "원본 PDF에서 배송현황 표의 상단 행들이 화면 하단 오버레이에 가려져 있어, 확인 가능한 행만 표로 복원했습니다.")
add_table(14, "배송현황(부분)", ["송장번호","고객명","발송일자","배송완료일"], [["TRK_004","EMMA","2025-10-04","NULL"]])
add_code(14, "SQL", "SELECT 송장번호, 고객명,\n       ______(배송완료일, '배송완료', '배송중') AS 배송상태\nFROM 배송현황;")

add_code(19, "SQL 스크립트", "\n".join([
"CREATE TABLE 주문로그 (LOG_ID NUMBER);",
"INSERT INTO 주문로그 VALUES (100);",
"INSERT INTO 주문로그 VALUES (200);",
"SAVEPOINT SP1;",
"INSERT INTO 주문로그 VALUES (300);",
"TRUNCATE TABLE 주문로그;",
"ROLLBACK TO SP1;",
"SELECT COUNT(*) FROM 주문로그;"
]))

add_table(20, "고객", ["고객ID","고객명","이메일"], [["C01","ALICE","alice@test.com"],["C02","BOB","NULL"],["C03","CHARLIE","charlie@test.com"],["C04","DAVID","NULL"]])

add_code(25, "메뉴 테이블(스키마)", "\n".join([
"(메뉴)",
"MENU_ID NUMBER PRIMARY KEY",
"NAME    VARCHAR2(10)",
"PRICE   NUMBER DEFAULT 0"
]))

add_table(28, "과제제출", ["제출ID","제출일시(DATE)"], [["S001","2025-06-15 00:00:00"],["S002","2025-06-15 15:30:00"],["S003","2025-06-15 23:59:59"],["S004","2025-06-16 00:00:00"]])
add_code(28, "SQL", "\n".join([
"SELECT COUNT(*)",
"FROM 과제제출",
"WHERE 제출일시 BETWEEN",
"TO_DATE('2025-06-15','YYYY-MM-DD')",
"AND TO_DATE('2025-06-15','YYYY-MM-DD') + 0.99999;"
]))

add_table(30, "제품", ["제품명","카테고리","재고"], [["P1","식품","5"],["P2","가전","20"],["P3","식품","10"],["P4","의류","5"]])
add_code(30, "SQL", "\n".join([
"SELECT COUNT(*)",
"FROM 제품",
"WHERE 카테고리 = '가전' OR 카테고리 = '식품' AND",
"재고 >= 10;"
]))

add_table(31, "사원", ["ENAME","DEPTNO","SALARY"], [["ALLEN","30","1600"],["WARD","30","1250"],["BLAKE","30","2850"],["KING","10","5000"]])
add_code(31, "SQL", "\n".join([
"SELECT ENAME, SALARY",
"FROM 사원",
"WHERE SALARY > ______ (SELECT SALARY",
"                       FROM 사원",
"                       WHERE DEPTNO = 30);"
]))

add_table(32, "월별매출", ["연도","월","매출액"], [["2024","12","800"],["2025","01","450"],["2025","02","550"]])
add_code(32, "SQL", "\n".join([
"(가)",
"SELECT 연도, 월, SUM(매출액)",
"  FROM 월별매출",
" GROUP BY ROLLUP(연도, 월);",
"",
"(나)",
"SELECT 연도, 월, SUM(매출액)",
"  FROM 월별매출",
" GROUP BY GROUPING SETS ( ______ );"
]))

add_table(35, "포인트", ["USERID","BONUS_A","BONUS_B"], [["User1","500","200"],["User2","300","100"],["User3","400","0"],["User4","NULL","100"],["User5","50","50"]])
add_code(35, "SQL", "SELECT SUM(BONUS_A + BONUS_B) + SUM(BONUS_A) AS TOTAL_POINT\nFROM 포인트;")

add_table(36, "A", ["ID"], [["10"],["20"]])
add_table(36, "B", ["ID"], [["10"],["30"]])
add_table(36, "C", ["ID"], [["10"],["40"]])
add_code(36, "SQL", "SELECT ID FROM A\nUNION ALL\nSELECT ID FROM B\nMINUS\nSELECT ID FROM C;")

add_code(38, "SQL 스크립트", "\n".join([
"CREATE TABLE 상품재고(",
" 상품ID NUMBER,",
" 상품명 VARCHAR2(100),",
" 재고수량 NUMBER DEFAULT 0",
");",
"INSERT INTO 상품재고 VALUES (101,'노트북',50);",
"INSERT INTO 상품재고 (상품ID, 상품명) VALUES (102,'마우스');",
"INSERT INTO 상품재고 VALUES (103,'키보드', NULL);",
"COMMIT;",
"SELECT SUM(재고수량) FROM 상품재고;"
]))

add_table(39, "메뉴", ["메뉴ID","메뉴명","상위메뉴ID"], [["1","음료","NULL"],["2","커피","1"],["3","아메리카노","2"]])
add_code(39, "SQL", "\n".join([
"SELECT 메뉴명, LEVEL",
"  FROM 메뉴",
" START WITH 메뉴명 = '아메리카노'",
" CONNECT BY ________;"
]))

add_code(41, "SQL 스크립트", "\n".join([
"CREATE TABLE EMP_LOG (EMPNO NUMBER);",
"INSERT INTO EMP_LOG VALUES (10);",
"SAVEPOINT A;",
"INSERT INTO EMP_LOG VALUES (20);",
"SAVEPOINT B;",
"INSERT INTO EMP_LOG VALUES (30);",
"ROLLBACK TO A;",
"INSERT INTO EMP_LOG VALUES (40);",
"COMMIT;",
"ROLLBACK;",
"SELECT COUNT(*) FROM EMP_LOG;"
]))

add_table(44, "결과", ["이름","수학","영어"], [["ALLEN","90","85"],["SMITH","80","95"]])
add_code(44, "SQL", "SELECT *\n  FROM 학생성적\n PIVOT ( MAX(점수) FOR ______ IN ('수학' AS 수학, '영어' AS 영어) );")

add_table(45, "직원", ["사원명","부서","급여"], [["KING","A","5000"],["SCOTT","A","3000"],["FORD","A","3000"],["SMITH","A","800"]])
add_table(45, "결과", ["사원명","부서","급여","순위"], [["KING","A","5000","1"],["SCOTT","A","3000","2"],["FORD","A","3000","2"],["SMITH","A","800","3"]])
add_code(45, "SQL", "SELECT 사원명, 부서, 급여,\n_______() OVER (PARTITION BY 부서 ORDER BY 급여 DESC) AS 순위\nFROM 직원;")

add_table(46, "사원", ["사번","이름","부서번호"], [["1001","SMITH","10"],["1002","ALLEN","20"]])

add_table(47, "일별매출", ["매출일자","매출액"], [["2024-01-01","100"],["2024-01-02","200"],["2024-01-03","300"],["2024-01-04","400"]])
add_code(47, "SQL", "\n".join([
"SELECT MAX(매출액) OVER () AS COL1,",
"       SUM(매출액) OVER (ORDER BY 매출일자) AS COL2,",
"       SUM(매출액) OVER (ORDER BY 매출일자",
"                         ROWS BETWEEN 1 PRECEDING AND 1 FOLLOWING) AS COL3",
"  FROM 일별매출;"
]))

add_code(48, "SQL 스크립트", "\n".join([
"CREATE TABLE 게시판 (",
" 글번호 NUMBER,",
" 제목 VARCHAR2(100),",
" 조회수 NUMBER",
");",
"INSERT INTO 게시판 (글번호, 제목) VALUES (1, '가입인사');",
"INSERT INTO 게시판 VALUES (2, '질문입니다', NULL);",
"INSERT INTO 게시판 VALUES (3, '공지사항', 10);",
"COMMIT;",
"SELECT SUM(조회수) FROM 게시판;"
]))

add_table(50, "상품", ["상품ID","상품명"], [["P001","상품A"],["P002","상품B"],["P003","상품C"],["P004","상품D"],["P005","상품E"]])
add_code(50, "SQL", "SELECT COUNT(*) AS CNT, MAX(상품ID) AS MAX_ID\nFROM 상품\nWHERE 1=0;")

# --- DOCX build ---
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
style.font.size = Pt(10)

title = doc.add_paragraph("59회 기출변형 문항 복원본")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(14)
title.runs[0].font.bold = False

doc.add_paragraph("")
doc.add_paragraph("주의: 이미지에 포함된 표는 가능한 범위에서 동일하게 표로 복원했습니다. 일부 문항은 원본 PDF에서 가려지거나 누락된 구간이 있어 주석을 달았습니다.")
doc.add_paragraph("")

circled = ["①","②","③","④","⑤","⑥"]

def add_code_block(document, code_text):
    t = document.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0,0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    tcPr.append(shd)
    # Clear existing content
    cell.text = ""
    # Remove the default empty paragraph
    if cell.paragraphs:
        p0 = cell.paragraphs[0]._element
        p0.getparent().remove(p0)
    for line in code_text.splitlines():
        para = cell.add_paragraph()
        run = para.add_run(line)
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

def add_table_block(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = str(h)
        for r in hdr_cells[j].paragraphs[0].runs:
            r.font.bold = False
    for row in rows:
        row_cells = table.add_row().cells
        for j, val in enumerate(row):
            row_cells[j].text = str(val)
    return table

for q in range(1, 51):
    stem = qa[q]["stem"]
    opts = qa[q]["opts"]

    para = doc.add_paragraph(f"{q}. {stem}")
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)

    for ex in qa[q]["extras"]:
        if ex["type"] == "note":
            n = doc.add_paragraph(f"[주석] {ex['note']}")
            n.paragraph_format.space_after = Pt(3)
        elif ex["type"] == "image":
            if ex.get("title"):
                doc.add_paragraph(f"[{ex['title']}]")
            doc.add_picture(ex["path"], width=Inches(ex.get("width_in", 6.0)))
        elif ex["type"] == "table":
            if ex.get("title"):
                doc.add_paragraph(f"[{ex['title']}]")
            add_table_block(doc, ex["headers"], ex["rows"])
        elif ex["type"] == "code":
            if ex.get("title"):
                doc.add_paragraph(f"[{ex['title']}]")
            add_code_block(doc, ex["code"])
        doc.add_paragraph("")

    for i, opt in enumerate(opts):
        prefix = circled[i] if i < len(circled) else f"({i+1})"
        doc.add_paragraph(f"{prefix} {opt}")
    doc.add_paragraph("")

out_docx="/mnt/data/59회기출변형_문항복원.docx"
doc.save(out_docx)

# Convert to PDF
tmp_profile = f"/tmp/lo_profile_{os.getpid()}"
os.makedirs(tmp_profile, exist_ok=True)
cmd = f"libreoffice -env:UserInstallation=file://{tmp_profile} --headless --convert-to pdf --outdir /mnt/data {shlex.quote(out_docx)}"
subprocess.run(cmd, shell=True, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
out_pdf="/mnt/data/59회기출변형_문항복원.pdf"

(out_docx, out_pdf)
